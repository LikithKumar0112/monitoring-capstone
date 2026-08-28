#!/usr/bin/env python3
# Minimal Markdown -> DOCX converter tailored to ARCHITECTURE.md.
# Handles: # ## ### headings, fenced code blocks (monospace, preserved),
# pipe tables, bullet/numbered lists, blockquotes, --- rules, and inline
# **bold** / `code`.
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

src, dst = sys.argv[1], sys.argv[2]
lines = open(src, encoding="utf-8").read().split("\n")

doc = Document()
# Base font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def add_inline(par, text):
    # split on **bold** and `code`, emit runs
    for tok in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            par.add_run(tok)


def add_code_block(code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    for i, cl in enumerate(code_lines):
        run = p.add_run(cl)
        run.font.name = "Consolas"
        run.font.size = Pt(7.5)   # small so wide ASCII diagrams fit the page
        if i != len(code_lines) - 1:
            run.add_break()


def add_table(rows):
    # rows: list of list-of-cell-strings; first row is header
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ""
            cell_par = cells[ci].paragraphs[0]
            add_inline(cell_par, txt.strip())
            for run in cell_par.runs:
                run.font.size = Pt(8.5)
                if ri == 0:
                    run.bold = True


def split_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return s.split("|")


i = 0
n = len(lines)
while i < n:
    line = lines[i]

    # fenced code block
    if line.strip().startswith("```"):
        i += 1
        buf = []
        while i < n and not lines[i].strip().startswith("```"):
            buf.append(lines[i])
            i += 1
        i += 1  # skip closing fence
        add_code_block(buf)
        continue

    # table: a line starting with | followed by a |---| separator
    if line.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
        rows = [split_row(line)]
        i += 2  # skip header + separator
        while i < n and lines[i].strip().startswith("|"):
            rows.append(split_row(lines[i]))
            i += 1
        add_table(rows)
        continue

    # headings
    m = re.match(r"^(#{1,4})\s+(.*)$", line)
    if m:
        level = len(m.group(1))
        doc.add_heading(m.group(2).strip(), level=min(level, 4))
        i += 1
        continue

    # horizontal rule
    if line.strip() == "---":
        doc.add_paragraph().add_run("").add_break()
        i += 1
        continue

    # blockquote
    if line.strip().startswith(">"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        content = re.sub(r"^\s*>\s?", "", line)
        r = p.add_run(content if content else " ")
        r.italic = True
        i += 1
        continue

    # numbered list
    m = re.match(r"^(\d+)\.\s+(.*)$", line)
    if m:
        p = doc.add_paragraph(style="List Number")
        add_inline(p, m.group(2))
        i += 1
        continue

    # bullet list
    m = re.match(r"^\s*[-*]\s+(.*)$", line)
    if m:
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, m.group(1))
        i += 1
        continue

    # blank
    if line.strip() == "":
        i += 1
        continue

    # normal paragraph
    p = doc.add_paragraph()
    add_inline(p, line)
    i += 1

doc.save(dst)
print("wrote", dst)
